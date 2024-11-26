import os
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Inches
import re

# 创建上一级目录中的文档保存目录
def create_document_directory():
    # 直接指定输出路径为 /app/document
    output_dir = "D:\work\python3.9\my_pc\document"
    os.makedirs(output_dir, exist_ok=True)  # 如果目录不存在，则创建
    return output_dir

# 生成 DOCX 文件
def generate_docx(content, image_path=None, filename="document.docx"):
    output_dir = create_document_directory()
    doc = Document()

    # 如果 filename 没有后缀，则添加 .docx 后缀
    if not filename.endswith(".docx"):
        filename += ".docx"

    # 解析文本文件内容
    # with open(input_text_file, 'r', encoding='utf-8') as file:
    #     content = file.read()

    # 提取各模块内容，使用 # 符号的数量来确定标题层级
    sections = []
    section_pattern = re.compile(r"(#+)\s*(.+)")
    lines = content.splitlines()
    current_section = {"heading": "", "content": "", "level": 0}

    for line in lines:
        match = section_pattern.match(line)
        if match:
            # 如果有正在构建的 section，将它加入 sections 列表
            if current_section["heading"] or current_section["content"]:
                sections.append(current_section)
                current_section = {"heading": "", "content": "", "level": 0}

            # 设置新的标题和层级
            level = len(match.group(1))  # # 符号的数量决定层级
            heading = match.group(2).strip()
            current_section = {"heading": heading, "content": "", "level": level}
        else:
            # 将非标题行加入当前 section 的内容
            current_section["content"] += line + "\n"

    # 添加最后一个 section
    if current_section["heading"] or current_section["content"]:
        sections.append(current_section)

    # 插入图片（可选）
    if image_path and os.path.exists(image_path):
        doc.add_picture(image_path, width=Inches(4))

    # 添加内容模块，根据层级设置标题格式
    for section in sections:
        if section["heading"]:
            heading = doc.add_heading(level=min(section["level"], 4))
            heading_run = heading.add_run(section["heading"])
            heading_run.bold = True
            heading_run.font.size = Pt(14 - section["level"])

        if section["content"].strip():
            paragraph = doc.add_paragraph(section["content"].strip())
            paragraph_format = paragraph.paragraph_format
            paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            paragraph_format.space_after = Pt(12)

    # 保存 DOCX 文件
    output_path = os.path.join(output_dir, filename)
    doc.save(output_path)



